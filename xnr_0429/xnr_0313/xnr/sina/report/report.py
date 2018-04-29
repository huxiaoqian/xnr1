#!/usr/bin/env python
#encoding: utf-8

from elasticsearch import Elasticsearch
import json
from screenshot import Screen
from docx import Document
import xlsxwriter
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

class Report:
	def __init__(self, id, username, password):
		self.ES_HOST = '219.224.134.213'
		self.ES_PORT = 9205
		self.INDEX_NAME = 'weibo_report_management'
		self.TYPE = 'report'
		self.es = Elasticsearch([{'host':self.ES_HOST,'port':self.ES_PORT}])
		self.results = []
		self.id = id
		self.username = username
		self.password = password

	def userList(self):
		self.results = []
		result = self.es.get(index=self.INDEX_NAME, doc_type=self.TYPE, id=self.id)['_source']
		event_name = result['event_name']
		report_time = result['report_time']
		report_type = result['report_type']
		xnr_user_no = result['xnr_user_no']

		weibo_list = json.loads(result['report_content'])['weibo_list']
		for each in weibo_list:
			text = each['text']
			timestamp = each['timestamp']
			try:
				user = each['nick_name']
			except:
				user = each['uid']
			uid = each['uid']
			mid = each['mid']

			dict = {'event_name':event_name, 'report_time':report_time, 'report_type':report_type,\
					'xnr_user_no':xnr_user_no, 'text':text, 'timestamp':timestamp, 'user':user,\
					'uid':uid, 'mid':mid}
			self.results.append(dict)
		return self.results

	def screen_shot(self, results):
		screen = Screen(self.username, self.password)
		for result in results:
			screen.screenShot(result['uid'], result['mid'])

	def save_excel(self):
		results = self.userList()
		if results:
			letters = "ABCDEFGHIJKLMN"
			self.screen_shot(results)
			file = xlsxwriter.Workbook(self.id+'.xlsx')
			table = file.add_worksheet()
			field = [each for each in results[0].keys()]
			field = '^&*'.join(field).replace('event_name',u'上报名称').replace('report_time',u'上报时间').replace('report_type',u'上报类型').replace('xnr_user_no',u'虚拟人').replace('text',u'文本内容').replace('user',u'发博用户').replace('timestamp',u'发博时间')+u"^&*截图"
			field = field.split('^&*')
			for a,b in enumerate(field):
				table.write(letters[a]+str(1),b)
			lists = []
			for result in results:
				list = [each for each in result.values()]
				lists.append(list)
			for i,k in enumerate(lists):
				table.insert_image(letters[len(k)]+str(16*i+2), results[0]['mid']+'.png', {'x_scale': 0.5, 'y_scale': 0.5})
				for c,d in enumerate(k):
					qq = letters[c]+str(i+2)
					table.write(qq,d)
			file.close()

	def save_word(self):
		results = self.userList()
		if results:
			self.screen_shot(results)
			document = Document()
			for result in results:
				str = json.dumps(result, ensure_ascii=False).replace('{','').replace('}','').replace('event_name',u'上报名称').replace('report_time',u'上报时间').replace('report_type',u'上报类型').replace('xnr_user_no',u'虚拟人').replace('text',u'文本内容').replace('user',u'发博用户').replace('timestamp',u'发博时间')
				document.add_paragraph(str)
				document.add_picture(result['mid']+'.png')
			document.add_page_break()
			document.save(self.id+'.docx')



if __name__ == '__main__':
	report = Report("WXNR0004_2154275107", "weiboxnr04@126.com", "xnr1234567")
	report.save_word()
	report.save_excel()



